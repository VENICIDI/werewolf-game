#!/usr/bin/env python3
"""
修复：清理论文重复引用 + 按首次出现重新编号 + 删除未引用文献
v2: 修复段落内多引用时按文本位置排序，而非数字排序
"""
import re
from docx import Document

SRC = r"C:\Users\dudu\Desktop\杨文杰-毕设初稿-重组版.docx"
doc = Document(SRC)

# ====== 找到参考文献起始位置 ======
ref_start = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("参考文献"):
        ref_start = i
        break

print(f"参考文献起始段落索引: {ref_start}")

# ====== 第一步：扫描正文引用（保留文本中的出现顺序）======
body_paras = []  # (para_index, [citation_nums_in_text_order])
for i in range(ref_start):
    text = doc.paragraphs[i].text
    # 使用 finditer 保持文本中出现顺序，每个数字只取第一次
    seen = set()
    nums_in_order = []
    for m in re.finditer(r'\[(\d+)\]', text):
        n = int(m.group(1))
        if n not in seen:
            seen.add(n)
            nums_in_order.append(n)
    if nums_in_order:
        body_paras.append((i, nums_in_order))

print(f"正文中含引用的段落数: {len(body_paras)}")
print("段落内引用顺序（前10段）:")
for pi, nums in body_paras[:15]:
    print(f"  段落{pi}: {nums}")

# ====== 第二步：按文本出现顺序编号 ======
old_to_new = {}
rank = 0
for pi, nums in body_paras:  # 已经按段落索引排序
    for n in nums:  # 段落内按文本顺序
        if n not in old_to_new:
            rank += 1
            old_to_new[n] = rank

print("\n旧 → 新 映射:")
for old in sorted(old_to_new):
    print(f"  [{old}] → [{old_to_new[old]}]")

# ====== 第三步：正文处理 ======
for pi, nums in body_paras:
    para = doc.paragraphs[pi]
    full_text = para.text

    # 找出所有 [数字] 标记及其位置
    matches = list(re.finditer(r'\[(\d+)\]', full_text))

    # 判断每个引用是否要保留（首次出现在本段）
    new_text = full_text
    for m in reversed(matches):
        num = int(m.group(1))
        start, end = m.span()

        # 检查这个引用是否首次出现在本段
        is_first = False
        for pi2, nums2 in body_paras:
            if pi2 == pi and num in nums2:
                is_first = True
                break
            elif pi2 < pi and num in nums2:
                # 在更早的段落已经出现 → 不是首次
                break

        if is_first and num in old_to_new:
            new_num = old_to_new[num]
            if new_num != num:
                new_text = new_text[:start] + f'[{new_num}]' + new_text[end:]
        else:
            # 删除此引用
            new_text = new_text[:start] + new_text[end:]

    if new_text != full_text:
        if para.runs:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = new_text
        else:
            para.text = new_text

# ====== 第四步：重构参考文献列表 ======
ref_list_start = ref_start + 1
ref_entries = []
for i in range(ref_list_start, len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    if text:
        m = re.match(r'\[(\d+)\]', text)
        if m:
            ref_entries.append((int(m.group(1)), i, text))

# 所有在引用中的编号
all_old = set(old_to_new.keys())

# 按新编号排序参考文献
ref_by_old = {old: (pi, text) for old, pi, text in ref_entries}
new_ref_order = sorted(old_to_new.items(), key=lambda x: x[1])

ref_para_indices = sorted(set(pi for _, pi, _ in ref_entries))

# 先清空所有参考文献段落
for pi in ref_para_indices:
    para = doc.paragraphs[pi]
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = ''
    else:
        para.text = ''

# 写入新参考文献
for idx, (old_num, new_num) in enumerate(new_ref_order):
    if old_num in ref_by_old:
        _, old_text = ref_by_old[old_num]
        new_text = re.sub(r'^\[\d+\]', f'[{new_num}]', old_text)
        target_pi = ref_para_indices[idx]
        para = doc.paragraphs[target_pi]
        if para.runs:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = new_text
        else:
            para.text = new_text

# 删除多余的参考文献段落
while len(new_ref_order) < len(ref_para_indices):
    pi = ref_para_indices[len(new_ref_order)]
    p_element = doc.paragraphs[pi]._element
    p_element.getparent().remove(p_element)
    ref_para_indices.pop()

# ====== 保存 ======
OUT = r"C:\Users\dudu\Desktop\杨文杰-毕设初稿-重组版-引用清理.docx"
doc.save(OUT)
print(f"\n✅ 已保存: {OUT}")
print(f"总引用数: {len(old_to_new)}，每个引用仅首次出现一次")
