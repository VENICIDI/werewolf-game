#!/usr/bin/env python3
"""验证引用清理结果"""
import re
from docx import Document

SRC = r"C:\Users\dudu\Desktop\杨文杰-毕设初稿-重组版-引用清理.docx"
doc = Document(SRC)

# 找到参考文献开始位置
ref_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("参考文献"):
        ref_start = i
        break

# 扫描正文引用
print("=" * 60)
print("正文引用检查（每个引用仅首次出现保留）")
print("=" * 60)

all_citations = {}  # new_num -> [para_indices]
for i in range(ref_start):
    text = doc.paragraphs[i].text
    found = re.findall(r'\[(\d+)\]', text)
    for n in found:
        n = int(n)
        if n not in all_citations:
            all_citations[n] = []
        all_citations[n].append(i)

# 检查重复
dups = {k: v for k, v in all_citations.items() if len(v) > 1}
if dups:
    print("\n❌ 仍有重复引用：")
    for k, v in sorted(dups.items()):
        print(f"  [{k}] 出现在段落: {v}")
else:
    print("\n✅ 无重复引用，每个引用只出现一次")

# 检查引用是否连续
cited_nums = sorted(all_citations.keys())
expected = list(range(1, len(cited_nums) + 1))
if cited_nums == expected:
    print(f"✅ 引用编号连续: {cited_nums}")
else:
    print(f"❌ 引用编号不连续: {cited_nums}, 期望: {expected}")

# 检查按段落出现顺序
print("\n正文引用出现顺序:")
for n in sorted(all_citations.keys(), key=lambda x: all_citations[x][0]):
    pis = all_citations[n]
    text = doc.paragraphs[pis[0]].text[:80]
    print(f"  [{n}] 段落{pis[0]}: {text}...")

# 扫描参考文献
print("\n" + "=" * 60)
print("参考文献列表")
print("=" * 60)
for i in range(ref_start + 1, len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    if text:
        m = re.match(r'\[(\d+)\]', text)
        if m:
            print(f"  [{m.group(1)}] {text[:80]}...")

print(f"\n总引用数: {len(cited_nums)}")
